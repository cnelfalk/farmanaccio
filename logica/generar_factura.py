# src/logica/generar_factura.py

import os
from docxtpl import DocxTemplate
from docx2pdf import convert
from tkinter.filedialog import asksaveasfilename
from tkinter import messagebox
from docx.shared import Pt

class FacturaGenerator:
    def __init__(self):
        # Ruta de la plantilla
        base = os.path.dirname(os.path.abspath(__file__))
        self.plantilla = os.path.join(base, "plantilla_factura.docx")

    def obtener_factura_y_detalles(self, cursor, factura_id):
        # Recupera datos principales, incluyendo tipoFactura
        cursor.execute("""
            SELECT facturaID, fechaEmision, horaEmision,
                   total_neto, total_bruto, descuento, tipoFactura
            FROM facturas
            WHERE facturaID = %s
        """, (factura_id,))
        factura = cursor.fetchone()
        detalles = []
        if factura:
            cursor.execute("""
                SELECT fd.prodID, p.nombre, fd.cantidad, fd.precioUnitario
                FROM factura_detalles fd
                JOIN productos p ON fd.prodID = p.prodID
                WHERE fd.facturaID = %s
            """, (factura_id,))
            filas = cursor.fetchall()
            for f in filas:
                f["subtotal"] = f["cantidad"] * f["precioUnitario"]
            detalles = filas
        return factura, detalles

    def insert_table_in_doc(self, doc, detalles):
        # Reemplaza el marcador de tabla en la plantilla
        for tabla in doc.tables:
            for row in tabla.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "%%tabla_placeholder%%" in para.text:
                            cell._tc.remove(para._element)
                            tbl = cell.add_table(rows=1, cols=4)
                            hdr = tbl.rows[0].cells
                            hdr[0].text, hdr[1].text = "Cantidad", "Producto"
                            hdr[2].text, hdr[3].text = "Precio Unit.", "Sub-total"
                            # Estilo de encabezados
                            for c in hdr:
                                for r in c.paragraphs[0].runs:
                                    r.font.name = "Helvetica"
                                    r.font.size = Pt(10)
                                    r.font.bold = True
                            # Filas de datos
                            for item in detalles:
                                rc = tbl.add_row().cells
                                rc[0].text = str(item["cantidad"])
                                rc[1].text = item["nombre"]
                                # Anteponer símbolo de peso $
                                rc[2].text = f"${item['precioUnitario']:.2f}"
                                rc[3].text = f"${item['subtotal']:.2f}"
                            return

    def generar_factura_con_transaccion(self, parent, conexion, factura_id, cliente: dict = None):
        """
        Genera la factura usando la conexión activa. Recibe opcionalmente un dict cliente:
          { 'nombre','apellido','cuit','iva' }
        Inserta el tipo de factura, el descuento real y rellena datos de cliente en la plantilla.
        """
        try:
            cursor = conexion.cursor(dictionary=True)
            datos, detalles = self.obtener_factura_y_detalles(cursor, factura_id)
            if not datos:
                messagebox.showerror("Error", "No se encontró la factura.", parent=parent)
                return False

            # Formateo del descuento
            dcto = float(datos["descuento"])
            descuento_str = "0%" if dcto == 0 else f"{dcto:.2f}%"

            # Círculos de IVA
            iva_val = (cliente or {}).get("iva", "").lower()
            circ = lambda cond: "●" if cond else "○"
            iva_ctx = {
                "ivaExento":      circ(iva_val == "exento"),
                "ivaMonotributo": circ(iva_val == "monotributo"),
                "ivaRespInsc":    circ(iva_val in ("resp. insc.", "responsable inscripto")),
                "ivaEventual":    circ(iva_val == "eventual"),
                "ivaConsFinal":   circ(iva_val == "cons. final"),
            }

            # Construcción de contexto con tipoFactura e IVA
            ctx = {
                "facturaID":     datos["facturaID"],
                "fecha":         datos["fechaEmision"],
                "hora":          datos["horaEmision"],
                "total_bruto":   datos["total_bruto"],
                "descuento":     descuento_str,
                "total_neto":    datos["total_neto"],
                "tipoFactura":   datos["tipoFactura"],
                "tabla_placeholder": "%%tabla_placeholder%%",
                **iva_ctx
            }

            # Datos de cliente
            if cliente:
                nombre_full = f"{cliente.get('nombre','')} {cliente.get('apellido','')}".strip()
                ctx["clienteNombre"]    = nombre_full
                ctx["clienteCUIT_CUIL"] = cliente.get("cuit","")
            else:
                ctx["clienteNombre"]    = ""
                ctx["clienteCUIT_CUIL"] = ""

            # Renderizado de la plantilla
            doc = DocxTemplate(self.plantilla)
            doc.render(ctx)
            self.insert_table_in_doc(doc, detalles)

            # Guardar .docx
            ruta = asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word","*.docx")],
                title="Guardar Factura"
            )
            if not ruta:
                messagebox.showwarning("Cancelado", "Guardado cancelado.", parent=parent)
                return False

            # Convertir a PDF
            doc.save(ruta)
            pdf = os.path.splitext(ruta)[0] + ".pdf"
            convert(ruta, pdf)
            os.remove(ruta)

            messagebox.showinfo("Éxito", f"Factura guardada en {pdf}", parent=parent)
            return True

        except Exception as ex:
            messagebox.showerror("Error Generando Factura", str(ex), parent=parent)
            return False

        finally:
            try:
                cursor.close()
            except:
                pass

    def generar_factura(self, parent, conexion, factura_id, cliente: dict = None):
        return self.generar_factura_con_transaccion(parent, conexion, factura_id, cliente)