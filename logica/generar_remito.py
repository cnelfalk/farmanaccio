import os
from datetime import datetime
from docxtpl import DocxTemplate
from docx2pdf import convert
from tkinter.filedialog import asksaveasfilename
from tkinter import messagebox
from docx.shared import Pt
from datos.conexion_bd import ConexionBD

class RemitoGenerator:
    def __init__(self):
        base = os.path.dirname(os.path.abspath(__file__))
        self.plantilla = os.path.join(base, "plantilla_remito.docx")
        self.clienteID = None  # MOD

    def insert_table_in_doc(self, doc, carrito):
        """
        Inserta la tabla de productos dentro de la plantilla del remito.
        Busca el marcador %%tabla_placeholder_remito%% y lo reemplaza
        por una tabla con columnas [Producto, Cantidad].
        """
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if "%%tabla_placeholder_remito%%" in para.text:
                            cell._tc.remove(para._element)
                            tbl = cell.add_table(rows=1, cols=2)
                            hdr = tbl.rows[0].cells
                            hdr[0].text, hdr[1].text = "Producto", "Cantidad"
                            for c in hdr:
                                for r in c.paragraphs[0].runs:
                                    r.font.name = "Helvetica"
                                    r.font.size = Pt(10)
                                    r.font.bold = True
                            for it in carrito:
                                rc = tbl.add_row().cells
                                rc[0].text = it.get("nombre", "")
                                rc[1].text = str(it.get("cantidad", ""))
                            return

    def insertar_en_bd(self, cliente, carrito, fecha_venc):
        """
        Inserta cabecera y detalle del remito en la base de datos.
        Ahora guarda también la dirección del cliente.
        Retorna el remitoID (auto-increment) o None si falla.
        """
        try:
            cnx = ConexionBD.obtener_conexion()
            cursor = cnx.cursor(dictionary=True)
            cursor.execute("USE farmanaccio_db")

            # Obtener clienteID por cuil-cuit
            cursor.execute(
                "SELECT clienteID FROM clientes WHERE `cuil-cuit` = %s",
                (cliente.get("cuit"),)
            )
            fila = cursor.fetchone()
            cid = fila["clienteID"] if fila else None
            # ➡️ MOD: guardamos el ID en el atributo
            self.clienteID = cid

            sql = """
                INSERT INTO Remito
                  (clienteID, cuit_cuil, ivaEstado, direccionCliente, fechaInicio, vencimientoRemito)
                VALUES (%s, %s, %s, %s, %s, %s)
            """
            hoy = datetime.now().date().isoformat()
            venc = fecha_venc.isoformat() if fecha_venc else None
            direccion = cliente.get("direccion", "")
            cursor.execute(sql, (
                cid,
                cliente.get("cuit"),
                cliente.get("iva"),
                direccion,
                hoy,
                venc
            ))
            remito_id = cursor.lastrowid

            for it in carrito:
                cursor.execute(
                    "INSERT INTO RemitoDetalle (remitoID, prodID, cantidad) VALUES (%s, %s, %s)",
                    (remito_id, it.get("prodID"), it.get("cantidad"))
                )

            cnx.commit()
            cursor.close()
            cnx.close()
            return remito_id

        except Exception as ex:
            try:
                cnx.rollback()
            except:
                pass
            messagebox.showerror("Error BD Remito", str(ex))
            return None

    def generar_remito_con_transaccion(self, parent, cliente, carrito, fecha_vencimiento=None):
        """
        Genera el remito: lo inserta en BD (con dirección),
        luego renderiza la plantilla y guarda .docx/.pdf.
        """
        try:
            iva_val = cliente.get("iva", "").lower()
            circ = lambda cond: "●" if cond else "○"
            iva_ctx = {
                "ivaExento":    circ(iva_val == "exento"),
                "ivaMonotributo": circ(iva_val == "monotributo"),
                "ivaRespInsc":  circ(iva_val in ("resp. insc.", "responsable inscripto")),
                "ivaEventual":  circ(iva_val == "eventual"),
                "ivaConsFinal": circ(iva_val == "cons. final")
            }

            remito_id = self.insertar_en_bd(cliente, carrito, fecha_vencimiento)
            if remito_id is None:
                return False

            ctx = {
                **iva_ctx,
                "remitoID":          remito_id,
                "fechaInicioRemito": datetime.now().date().isoformat(),
                "horaInicioRemito":  datetime.now().time().strftime("%H:%M:%S"),
                "fechaVencRemito":   fecha_vencimiento.isoformat() if fecha_vencimiento else "",
                "clienteNombre":     f"{cliente.get('nombre','')} {cliente.get('apellido','')}".strip(),
                "clienteCUIT_CUIL":  cliente.get("cuit",""),
                "clienteDireccion":  cliente.get("direccion",""),
                "%%tabla_placeholder_remito%%": "%%tabla_placeholder_remito%%",
                # ➡️ MOD: ponemos aquí el ID
                "clienteID":         self.clienteID
            }

            doc = DocxTemplate(self.plantilla)
            doc.render(ctx)
            self.insert_table_in_doc(doc, carrito)

            ruta = asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word","*.docx")],
                title="Guardar Remito"
            )
            if not ruta:
                messagebox.showwarning("Cancelado", "Guardado cancelado.", parent=parent)
                return False

            doc.save(ruta)
            pdf = os.path.splitext(ruta)[0] + ".pdf"
            convert(ruta, pdf)
            os.remove(ruta)
            messagebox.showinfo("Éxito", f"Remito guardado en {pdf}", parent=parent)
            return True

        except Exception as ex:
            messagebox.showerror("Error Remito", str(ex), parent=parent)
            return False

    def generar_remito(self, parent, cliente, carrito, fecha_vencimiento=None):
        return self.generar_remito_con_transaccion(
            parent, cliente, carrito, fecha_vencimiento
        )